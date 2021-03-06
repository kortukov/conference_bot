from classes import (
    BaseEvent,
    SimpleEvent,
    FullEvent,
    Food,
    Other,
    Plenary,
    Research,
    Young,
    OtherFull,
    Talk,
)
from datetime import datetime
import time
from docx import Document
import re
import sys
import pickle
import jsonpickle
import json

PROGRAM_PATH = "../Program_example.docx"


def get_timestamp(date, hours, minutes):
    ''' date is day(int), hours and minutes are also (int)'''
    dt = datetime(year=2019, month=9, day=date, hour=hours, minute=minutes)
    return time.mktime(dt.timetuple())


class ProgramParser:
    def __init__(self):
        pass

    @staticmethod
    def get_place_id(string):
        places = {
            "Sokolniki": 0,
            "Сокольники": 0,
            "Passage": 1,
            "Пассаж": 1,
            "Москва": 2,
            "Moskva": 2,
            "Okhotny": 3,
            "Охотный": 3,
            "Krymsky": 4,
            "Крымский": 4,
            "Vorobiovy": 5,
            "Воробьевы": 5,
            "Воробьёвы": 5,
            "Arbat": 6,
            "Арбат": 6,
            "Krasnye": 7,
            "Красные": 7,
            "Ostozhenka": 8,
            "Остоженка": 8,
            "Chistye": 9,
            "Чистые": 9,
            "Polyanka": 10,
            "Полянка": 10,
            "Maroseika": 11,
            "Маросейка": 11,
        }

        string = re.sub('[«»“”]', '', string)
        for word in string.split(" "):

            if word in places:
                return places[word]
        return -1

    def parse_program(self, program_path):
        document = Document(program_path)
        doc_length = len(document.paragraphs)

        event_list = []
        full_event_counter = 0
        talks_counter = 0
        for i in range(1, doc_length):
            text = document.paragraphs[i].text

            if "Понедельник" in text:
                if 'Понедельник, 23 сентября 2019 г' in text:
                    continue
                current_date = 23
            elif "Вторник" in text:
                if 'Вторник, 24 сентября 2019 г' in text:
                    continue
                current_date = 24
            else:
                continue
            if "Poster" in text:
                continue
            if document.paragraphs[i + 1].text == "":  # leaving tables out
                continue

            current_time = document.paragraphs[i].text.split()[3]
            current_place = self.get_place_id(document.paragraphs[i + 1].text)
            # getting timestamps from begin and end times
            begin = current_time.split('-')[0]
            end = current_time.split('-')[1]

            ts_begin = get_timestamp(
                current_date, int(begin.split(':')[0]), int(begin.split(':')[1])
            )
            ts_end = get_timestamp(current_date, int(end.split(':')[0]), int(end.split(':')[1]))

            title_ru = document.paragraphs[i + 2].text.split("//")[0]
            if "//" in document.paragraphs[i + 2].text:
                title_en = document.paragraphs[i + 2].text.split("//")[1]
                if title_en[0] == " ":
                    title_en = title_en[1:]  # cleaning one space in the beginning
            else:
                title_en = title_ru
            title_ru = title_ru.rstrip()
            # now deciding which object to create
            if any(substring in title_en for substring in ["Lunch", "Reception", "Coffee"]):
                event = Food(ts_begin, ts_end, current_place, title_ru, title_en)

            elif "Plenary" in title_en:
                event = Plenary(
                    ts_begin, ts_end, current_place, title_ru, title_en, full_event_counter
                )
                full_event_counter += 1

            elif "Research" in title_en:
                event = Research(
                    ts_begin, ts_end, current_place, title_ru, title_en, full_event_counter
                )
                full_event_counter += 1

            elif "Student" in title_en:
                event = Young(
                    ts_begin, ts_end, current_place, title_ru, title_en, full_event_counter
                )
                full_event_counter += 1

            elif "Awards" in title_en:
                event = Other(
                    ts_begin, ts_end, current_place, title_ru, title_en, full_event_counter
                )
                full_event_counter += 1

            else:
                # now checking if list of talks is needed
                j = 1
                list_needed = 0
                while i + 2 + j < doc_length and not "//" in document.paragraphs[i + 2 + j].text:
                    for run in document.paragraphs[i + 2 + j].runs:
                        if run.bold:
                            list_needed = 1
                    j = j + 1
                if list_needed:
                    event = OtherFull(
                        ts_begin, ts_end, current_place, title_ru, title_en, full_event_counter
                    )
                    full_event_counter += 1
                else:
                    event = Other(
                        ts_begin, ts_end, current_place, title_ru, title_en, full_event_counter
                    )
                    full_event_counter += 1

            if isinstance(event, FullEvent) or isinstance(event, Other):
                # filling the list of talks
                talks_list = []
                j = 1
                while i + 2 + j < doc_length and not "//" in document.paragraphs[i + 2 + j].text:
                    talks_list = talks_list + document.paragraphs[i + 2 + j].runs
                    j = j + 1

                j = 1
                while j < len(talks_list):  # clearing list of talks
                    talks_list[j].text = talks_list[j].text.rstrip()
                    if talks_list[j].text == '':
                        del talks_list[j]
                    j = j + 1

                j = 1
                while j < len(talks_list):  # looking for talks and authors
                    if talks_list[j].bold:
                        current_talk = talks_list[j].text
                        current_talk = current_talk.strip('\n')
                        current_authors = ""
                        current_speaker = ""
                        k = 1
                        while (j + k < len(talks_list)) and (not talks_list[j + k].bold):
                            if talks_list[j + k].underline:
                                current_speaker = talks_list[j + k].text
                            if talks_list[j + k].font.superscript:
                                k = k + 1
                                continue
                            current_authors = current_authors + talks_list[j + k].text.strip('\n')

                            k = k + 1

                        event.talks_list.append(
                            Talk(
                                title=current_talk,
                                authors=current_authors,
                                speaker=current_speaker,
                                hall=current_place,
                                talk_number=talks_counter,
                            )
                        )
                        talks_counter += 1

                    j = j + 1

                if not event.talks_list:
                    talk = Talk(
                        title=event.title_ru,
                        authors='',
                        speaker='',
                        hall=current_place,
                        talk_number=talks_counter,
                    )

                    event.talks_list.append(talk)
                    talks_counter += 1

            if isinstance(event, Other):  # Looking for descriptio
                description_lines = []
                description = ""
                j = 1
                while i + 2 + j < doc_length and not "//" in document.paragraphs[i + 2 + j].text:
                    description_lines = description_lines + document.paragraphs[i + 2 + j].runs
                    j = j + 1
                for t in description_lines:
                    t.text = t.text.strip()
                    t.text = t.text.strip('\n')
                    if t.text != '':
                        description += t.text + '\n'
                if description != "":
                    event.description = description

            event_list.append(event)

        for event in event_list:
            if isinstance(event, FullEvent):
                event.calculate_talk_times()

        return event_list[:-1]


if __name__ == '__main__':

    if len(sys.argv) == 1:
        mode = 'print'
    elif len(sys.argv) == 2 and (sys.argv[1] == 'json' or sys.argv[1] == 'pickle'):
        mode = sys.argv[1]
    else:
        sys.exit("Usage: python3 classes.py [pickle / json]\nDefault is printing.")

    parser = ProgramParser()
    event_list = parser.parse_program(PROGRAM_PATH)

    if mode == 'print':
        for event in event_list:
            print(event.event_type)
            if isinstance(event, FullEvent):
                print(event.full_str_ru())
            else:
                print(event)
            print("================")
    elif mode == 'pickle':
        with open('../event_list.pickle', 'wb') as f:
            pickle.dump(event_list, f, pickle.HIGHEST_PROTOCOL)
    elif mode == 'json':

        with open('../event_list.json', 'w') as f:
            for event in event_list:
                f.write(jsonpickle.encode(event))
                f.write('\n')
