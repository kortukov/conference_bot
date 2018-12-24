from docx import Document
import re
class BaseEvent:

    def __init__(self, timestamp, place_id, event_type, title_ru, title_en):
        self.timestamp = timestamp
        self.place_id = place_id
        self.event_type = event_type
        self.title_ru = title_ru
        self.title_en = title_en
    def __str__(self):
        return "" + str(self.timestamp) + " " + str(self.place_id) + " " + str(self.event_type) + " " + str(self.title_ru) + " " + str(self.title_en)

    def getEventType(self):
        return self.event_type

#Events with no description: Food, Registration, other without description - only a name for them
class SimpleEvent(BaseEvent):
    def __init__(self, timestamp, place_id, event_type, title_ru, title_en):
        super().__init__(timestamp, place_id, event_type, title_ru, title_en)
    def __str__(self):
        return super().__str__()

class Food(SimpleEvent):
    def __init__(self, timestamp, place_id, title_ru, title_en):
         super().__init__(timestamp, place_id, "Food", title_ru, title_en)

class Other(SimpleEvent):
    def __init__(self, timestamp, place_id, title_ru, title_en):
         super().__init__(timestamp, place_id, "Other", title_ru, title_en)

#Events with description: Plenary session, research session, Young researchers, other with description
class FullEvent(BaseEvent):
    
    def __init__(self, timestamp, place_id, event_type, title_ru, title_en):
        self.sublist = []
        super().__init__(timestamp, place_id, event_type, title_ru, title_en)
    
    def __str__(self):
        result = ""
        result = super().__str__() + '\n' 
        for element in self.sublist:
            result = result + "Title:\n\t" + element[0] + '\n'
            result = result + "Authors\n\t" + element[1] + '\n'
        return result

class Plenary(FullEvent):
    def __init__(self, timestamp, place_id, title_ru, title_entle):
         super().__init__(timestamp, place_id, "Plenary", title_ru, title_en)

class Research(FullEvent):

    def __init__(self, timestamp, place_id, title_ru, title_en):
         super().__init__(timestamp, place_id, "Research", title_ru, title_en)
        

class Young(FullEvent):
    def __init__(self, timestamp, place_id, title_ru, title_en):
         super().__init__(timestamp, place_id, "Young researchers", title_ru, title_en)

class OtherFull(FullEvent):
    def __init__(self, timestamp, place_id, title_ru, title_en):
         super().__init__(timestamp, place_id, "Other full events", title_ru, title_en)        
         

def get_place_id(string):
    places = {
    "Sokolniki" : 0,
    "Сокольники" : 0,
    "Passage" : 1,
    "Пассаж" : 1,
    "Москва" : 2,
    "Moskva" : 2,
    "Okhotny" : 3,
    "Охотный" : 3,
    "Krymsky" : 4,
    "Крымский" : 4,
    "Vorobiovy" : 5,
    "Воробьевы" : 5,
    "Воробьёвы" : 5,
    "Arbat" : 6,
    "Арбат" : 6,
    "Krasnye" : 7,
    "Красные": 7,
    "Ostozhenka" : 8,
    "Остоженка" : 8,
    "Chistye" : 9,
    "Чистые" : 9,
    "Polyanka" : 10,
    "Полянка" : 10,
    "Maroseika" : 11,
    "Маросейка" : 11,
    } 
    
    string = re.sub('[«»“”]', '', string)
    for word in string.split(" "):

        if word in places:
            return places[word]
    return -1


"""a = Research("25 September 12:00", 1 , "Первый доклад", "First talk in research")    
b = Young("25 September 12:00", 1, "Молодый специалисты","First talk in young researchers")
c = Food("25 September 14:30", 2, "Обед","Lunch")
l = [a,b,c]
for obj in l:
    print(obj)"""

document = Document("Program-2018_v4.docx")
doc_length = len(document.paragraphs)
for i in range(1,doc_length):
    text = document.paragraphs[i].text

    if "Понедельник" in text:
        current_date = 24
    elif "Вторник" in text:
        current_date = 25
    else:
        continue
    if "Poster" in text:
        continue
    if (document.paragraphs[i+1].text) == "": #leaving tables out
       continue

    current_time = document.paragraphs[i].text.split()[3]
    current_place = get_place_id(document.paragraphs[i+1].text)
    #timestamp function from current_date and current_time

    title_ru = document.paragraphs[i+2].text.split("//")[0]
    if "//" in document.paragraphs[i+2].text:
        title_en = document.paragraphs[i+2].text.split("//")[1]
        if title_en[0] == " ":
            title_en = title_en[1:] #cleaning one space in the beginning
    else:
        title_en = title_ru
    title_ru = title_ru.rstrip()
    #now deciding which object to create
    if any(substring in title_en for substring in ["Lunch", "Reception", "Coffee"]):
        event = Food(current_time, current_place, title_ru, title_en)

    elif "Plenary" in title_en:
        event = Plenary(current_time, current_place, title_ru, title_en)

    elif "Research" in title_en:
        event = Research(current_time, current_place, title_ru, title_en)

    elif "Student" in title_en:
        event = Young(current_time, current_place, title_ru, title_en)  

    elif "Awards" in title_en:
        event = Other(current_time, current_place, title_ru, title_en)
    
    else:
        #now checking if list of talks is needed
        j = 1
        list_needed = 0
        while not "//" in document.paragraphs[i+2+j].text:
            for run in  document.paragraphs[i+2+j].runs:
                if run.bold:
                    list_needed = 1
            j = j + 1
        if list_needed:
            event = OtherFull(current_time, current_place, title_ru, title_en)
        else:
            event = Other(current_time, current_place, title_ru, title_en)

    if isinstance(event, FullEvent):
        #filling the list of talks
        talks_list = []
        j = 1
        while not "//" in document.paragraphs[i+2+j].text:
            talks_list = talks_list + document.paragraphs[i+2+j].runs
            j = j + 1

        j = 1    
        while j < len(talks_list): #clearing list of talks
            talks_list[j].text = talks_list[j].text.rstrip()
            if talks_list[j].text == '':
                del(talks_list[j])
            j = j + 1

        j = 1    
        while j < len(talks_list): #looking for talks and authors
            if talks_list[j].bold:
                current_talk = talks_list[j].text
                current_talk = current_talk.strip('\n')
                current_authors = ""
                k = 1
                while (j+k < len(talks_list)) and (not talks_list[j+k].bold): #using lazy logic
                    current_authors = current_authors +  talks_list[j+k].text.strip('\n')

                    k = k + 1
                event.sublist.append((current_talk, current_authors))
            j = j + 1

        """for elem in event.sublist:
            print("---")
            print(elem[0])
            print("Authors are:")
            print(elem[1])"""
        #print(event)
        #print("=========")        
             
    print(event)
    print("================")
